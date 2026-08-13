import sys
import subprocess
import os
import shutil
import importlib.metadata
import time

# External Dataset Location
EXTERNAL_DATASET_PATH = 'K:/ImageDataset'

# ==========================================
# 0. Environment Bootstrap & Dependencies
# ==========================================
def install_dependencies():
    """Installs required packages if they are not already present in the environment."""
    # PyTorch does not support Python 3.14+ yet.
    if sys.version_info >= (3, 14):
        print(f"[ERROR] Python {sys.version_info.major}.{sys.version_info.minor} is not supported by PyTorch yet.")
        print("[ERROR] Please use Python 3.10, 3.11, or 3.12.")
        sys.exit(1)

    required_packages = [
        "Pillow", "numpy", "matplotlib", "seaborn",
        "scikit-learn", "pandas", "boto3", "psutil", "python-docx"
    ]

    def is_installed(package_name):
        try:
            importlib.metadata.version(package_name)
            return True
        except importlib.metadata.PackageNotFoundError:
            return False

    missing_packages = [pkg for pkg in required_packages if not is_installed(pkg)]
    
    def is_nvidia_gpu_present():
        try:
            subprocess.check_output(["nvidia-smi"], stderr=subprocess.STDOUT)
            return True
        except Exception:
            return False

    gpu_present = is_nvidia_gpu_present()
    
    try:
        import torch
        import torchvision
        torch_ok = True
        cuda_ok = torch.cuda.is_available()
    except ImportError:
        torch_ok = False
        cuda_ok = False

    # Check for missing torch packages specifically
    torch_packages = ["torch", "torchvision", "torchaudio"]
    missing_torch = [pkg for pkg in torch_packages if not is_installed(pkg)]
    
    # We need to install general dependencies if any are missing
    if missing_packages:
        print(f"[INFO] Environment check: Missing dependencies {missing_packages}. Installing...")
        subprocess.check_call([sys.executable, "-m", "pip", "install", *missing_packages])

    # We need to install/reinstall torch if it's missing or if GPU is present but CUDA is not available
    if missing_torch or (gpu_present and not cuda_ok):
        install_cmd = [sys.executable, "-m", "pip", "install"]
        
        if gpu_present and not cuda_ok:
            if torch_ok:
                print("[WARN] NVIDIA GPU detected but PyTorch is CPU-only. Reinstalling PyTorch with CUDA support...")
                install_cmd.extend(["torch", "torchvision", "torchaudio", "--force-reinstall"])
            else:
                print("[INFO] NVIDIA GPU detected. Installing PyTorch with CUDA support...")
                install_cmd.extend(["torch", "torchvision", "torchaudio"])
        else:
            print(f"[INFO] Environment check: Missing PyTorch dependencies {missing_torch}. Installing...")
            install_cmd.extend(missing_torch)
            
        install_cmd.extend(["--index-url", "https://download.pytorch.org/whl/cu121"])
        subprocess.check_call(install_cmd)
        print("[SUCCESS] PyTorch dependencies handled successfully.\n")
    
    if not missing_packages and not missing_torch and not (gpu_present and not cuda_ok):
        # All good, but let's print something if we were verbose before
        # print("[INFO] Environment check: All dependencies are present.")
        pass


# Environment check logic moved to __main__ block

# Late imports ensure dependencies are installed before they are called
import json
import math
import random
import argparse
import platform
import psutil
import numpy as np
import pandas as pd
import matplotlib.pyplot as plt
import seaborn as sns
from PIL import Image
from matplotlib.backends.backend_pdf import PdfPages

import torch
import torch.nn as nn
import torch.optim as optim
from torch.utils.data import DataLoader
from torchvision import models, datasets, transforms

from sklearn.metrics import (
    accuracy_score, precision_score, recall_score, f1_score,
    confusion_matrix, classification_report,
    roc_curve, auc,
    precision_recall_curve, average_precision_score,
    matthews_corrcoef, balanced_accuracy_score, brier_score_loss
)

# ==========================================
# 1. Configuration & Path Management
# ==========================================
# Uses environment variables first, falls back to SageMaker standards, then defaults to standard local directories.
BASE_DIR = os.environ.get("BASE_DIR", os.getcwd())

# Datasets: Default to SageMaker's training channel if it exists, otherwise use local './ImageDataset'
DEFAULT_DATASET_PATH = os.environ.get("SM_CHANNEL_TRAINING", os.path.join(BASE_DIR, EXTERNAL_DATASET_PATH))
DATASET_PATH = os.environ.get("DATASET_DIR", DEFAULT_DATASET_PATH)

# Logs/Metrics: Default to './logs'
LOGS_PATH = os.environ.get("LOGS_DIR", os.path.join(BASE_DIR, 'logs'))
os.makedirs(LOGS_PATH, exist_ok=True)

def initialize_logging():
    """Initializes logging to both console and a master log file (appends each run)."""
    log_file = os.path.join(LOGS_PATH, "pipeline_execution.log")
    
    class Tee:
        def __init__(self, terminal, stream):
            self.terminal = terminal
            self.stream = stream
        def write(self, message):
            self.terminal.write(message)
            self.stream.write(message)
            self.stream.flush()
        def flush(self):
            self.terminal.flush()
            self.stream.flush()
        def __getattr__(self, attr):
            return getattr(self.terminal, attr)

    try:
        # Open in append mode to keep history of all executions
        log_stream = open(log_file, "a", encoding="utf-8")
        log_stream.write(f"\n{'='*70}\n")
        log_stream.write(f"SESSION START: {time.strftime('%Y-%m-%d %H:%M:%S')}\n")
        log_stream.write(f"Command: {' '.join(sys.argv)}\n")
        log_stream.write(f"{'='*70}\n\n")
        
        sys.stdout = Tee(sys.stdout, log_stream)
        sys.stderr = Tee(sys.stderr, log_stream)
        print(f"[INFO] Session logging started (appending to {log_file})")
    except Exception as e:
        print(f"[WARN] Could not initialize log file: {e}")

# initialize_logging() # Moved to __main__ block to avoid multiple logs from data workers

# Results/Outputs: Only PDF and CSV should be stored here.
DEFAULT_RESULTS_PATH = os.environ.get("SM_MODEL_DIR", os.path.join(BASE_DIR, 'results'))
RESULTS_PATH = os.environ.get("RESULTS_DIR", DEFAULT_RESULTS_PATH)
os.makedirs(RESULTS_PATH, exist_ok=True)

# Artifacts & Models: Intermediate files (PNG, JSON) and model weights (.pth)
# are stored in the logs directory to keep the results folder clean.
ARTIFACTS_PATH = os.path.join(LOGS_PATH, 'artifacts')
os.makedirs(ARTIFACTS_PATH, exist_ok=True)

MODELS_PATH = os.path.join(LOGS_PATH, 'models')
os.makedirs(MODELS_PATH, exist_ok=True)

IMG_SIZE = int(os.environ.get("IMG_SIZE", 150))

def verify_directories():
    """Ensures that all required directories exist and are accessible."""
    print("[INFO] Path verification in progress...")
    
    # 1. Dataset Path (Essential for training/eval)
    if not os.path.exists(DATASET_PATH):
        print(f"[WARN] DATASET_PATH does not exist: {DATASET_PATH}")
        print("       (Execution may fail if the current mode requires local data)")
    elif not os.path.isdir(DATASET_PATH):
        print(f"[ERROR] DATASET_PATH exists but is not a directory: {DATASET_PATH}")
    else:
        print(f"[SUCCESS] DATASET_PATH verified: {DATASET_PATH}")

    # 2. Logs and Results (Must be writable)
    for path_name, path_val in [("LOGS_PATH", LOGS_PATH), ("RESULTS_PATH", RESULTS_PATH)]:
        try:
            os.makedirs(path_val, exist_ok=True)
            # Check writability
            test_file = os.path.join(path_val, '.write_test')
            with open(test_file, 'w') as f:
                f.write('test')
            os.remove(test_file)
            print(f"[SUCCESS] {path_name} verified: {path_val}")
        except Exception as e:
            print(f"[ERROR] {path_name} cannot be initialized: {path_val}")
            print(f"        Details: {e}")
            sys.exit(1)

def clear_directories():
    """Deletes all files in the results, models, and artifacts directories."""
    targets = [
        ("Results", RESULTS_PATH),
        ("Models", MODELS_PATH),
        ("Artifacts", ARTIFACTS_PATH)
    ]
    
    print("\n--- Clearing Pipeline Directories ---")
    for name, path in targets:
        if os.path.exists(path):
            items_removed = 0
            for item in os.listdir(path):
                item_path = os.path.join(path, item)
                try:
                    if os.path.isfile(item_path) or os.path.islink(item_path):
                        os.unlink(item_path)
                        items_removed += 1
                    elif os.path.isdir(item_path):
                        shutil.rmtree(item_path)
                        items_removed += 1
                except Exception as e:
                    print(f"[ERROR] Failed to delete {item_path}. Reason: {e}")
            print(f"[SUCCESS] {name} directory cleared ({items_removed} items removed).")
        else:
            print(f"[INFO] {name} directory does not exist: {path}")
    print("--------------------------------------\n")

# Directory verification logic moved to __main__ block
device = torch.device("cuda" if torch.cuda.is_available() else "cpu")


# ==========================================
# Hardware Diagnostic Report
# ==========================================
def print_hardware_report():
    print("\n--- Hardware Diagnostic Report ---")

    # CPU Information
    cpu_name = platform.processor() or "Unknown CPU Architecture"
    print(f"CPU Name: {cpu_name}")

    # System Memory (RAM)
    sys_mem = psutil.virtual_memory()
    total_mem_gb = sys_mem.total / (1024 ** 3)
    print(f"System Memory (RAM): {total_mem_gb:.2f} GB")

    # Target Device Mapping
    device_type = "cuda" if torch.cuda.is_available() else "cpu"
    print(f"Target Compute Device: {device_type.upper()}")

    # GPU Information
    if torch.cuda.is_available():
        gpu_count = torch.cuda.device_count()
        print(f"Number of GPUs: {gpu_count}")
        for i in range(gpu_count):
            gpu_name = torch.cuda.get_device_name(i)
            vram_gb = torch.cuda.get_device_properties(i).total_memory / (1024 ** 3)
            print(f"  - GPU {i}: {gpu_name} ({vram_gb:.2f} GB VRAM)")
    else:
        print("Number of GPUs: 0")
        print("  - No GPU detected. Execution will fall back to CPU.")

    print("--------------------------------------\n")


# ==========================================
# 2. Model Architecture
# ==========================================
AVAILABLE_MODELS = ["DenseNet121", "ResNet50", "EfficientNetB0"]

def list_available_models():
    """Prints the list of supported model architectures."""
    print("\n--- Available Model Architectures ---")
    for model in AVAILABLE_MODELS:
        print(f"  - {model}")
    print("--------------------------------------\n")

def build_model(model_name, pretrained=True):
    print(f"Building PyTorch model: {model_name}...")
    if model_name == "ResNet50":
        weights = models.ResNet50_Weights.DEFAULT if pretrained else None
        model = models.resnet50(weights=weights)
        in_features = model.fc.in_features
        model.fc = nn.Sequential(
            nn.Linear(in_features, 512), nn.ReLU(), nn.Dropout(0.5), nn.Linear(512, 1)
        )
    elif model_name == "DenseNet121":
        weights = models.DenseNet121_Weights.DEFAULT if pretrained else None
        model = models.densenet121(weights=weights)
        in_features = model.classifier.in_features
        model.classifier = nn.Sequential(
            nn.Linear(in_features, 512), nn.ReLU(), nn.Dropout(0.5), nn.Linear(512, 1)
        )
    elif model_name == "EfficientNetB0":
        weights = models.EfficientNet_B0_Weights.DEFAULT if pretrained else None
        model = models.efficientnet_b0(weights=weights)
        in_features = model.classifier[1].in_features
        model.classifier[1] = nn.Sequential(
            nn.Linear(in_features, 512), nn.ReLU(), nn.Dropout(0.5), nn.Linear(512, 1)
        )
    else:
        raise ValueError(f"Unknown model name: {model_name}")

    if pretrained:
        for param in model.parameters():
            param.requires_grad = False
        if model_name == "ResNet50":
            for param in model.fc.parameters(): param.requires_grad = True
        elif model_name == "DenseNet121":
            for param in model.classifier.parameters(): param.requires_grad = True
        elif model_name == "EfficientNetB0":
            for param in model.classifier.parameters(): param.requires_grad = True
    return model


# ==========================================
# 3. Metrics & Reporting
# ==========================================
class MetricsManager:
    def __init__(self):
        self.start_time = 0
        self.end_time = 0

    def start_timer(self):
        self.start_time = time.time()

    def stop_timer(self):
        self.end_time = time.time()
        return self.end_time - self.start_time

    def calculate_metrics(self, y_true, y_pred_probs):
        y_pred = [1 if p > 0.5 else 0 for p in y_pred_probs]
        return {
            "accuracy": accuracy_score(y_true, y_pred),
            "precision": precision_score(y_true, y_pred, zero_division=0),
            "recall": recall_score(y_true, y_pred, zero_division=0),
            "f1_score": f1_score(y_true, y_pred, zero_division=0),
            "balanced_accuracy": balanced_accuracy_score(y_true, y_pred),
            "mcc": matthews_corrcoef(y_true, y_pred),
            "brier_score": brier_score_loss(y_true, y_pred_probs),
            "confusion_matrix": confusion_matrix(y_true, y_pred).tolist(),
            "report": classification_report(y_true, y_pred, output_dict=True, zero_division=0)
        }

    def plot_confusion_matrix(self, cm, labels=['Normal', 'Anomaly'], title="Confusion Matrix", epochs=None):
        fig = plt.figure(figsize=(6, 4))
        sns.heatmap(cm, annot=True, fmt='d', cmap='Blues', xticklabels=labels, yticklabels=labels)
        plt.xlabel('Predicted')
        plt.ylabel('Actual')
        full_title = f"{title} ({epochs} Epochs)" if epochs is not None else title
        plt.title(full_title)
        plt.tight_layout()
        return fig

    def plot_roc_curve(self, y_true, y_probs, dataset_name="unknown", model_name="unknown", save_dir=ARTIFACTS_PATH, epochs=None):
        fpr, tpr, _ = roc_curve(y_true, y_probs)
        roc_auc = auc(fpr, tpr)
        roc_data = {"fpr": fpr.tolist(), "tpr": tpr.tolist(), "auc": float(roc_auc)}

        json_path = os.path.join(save_dir, f"roc_data_{dataset_name}_{model_name}.json")
        with open(json_path, 'w') as f:
            json.dump(roc_data, f)

        fig = plt.figure(figsize=(8, 6))
        plt.plot(fpr, tpr, color='darkorange', lw=2, label=f'ROC curve (area = {roc_auc:.4f})')
        plt.plot([0, 1], [0, 1], color='navy', lw=2, linestyle='--', label='Random classifier')
        plt.xlim([0.0, 1.0])
        plt.ylim([0.0, 1.05])
        plt.xlabel('False Positive Rate')
        plt.ylabel('True Positive Rate')
        epoch_str = f", {epochs} Epochs" if epochs is not None else ""
        plt.title(f'ROC Curve - {dataset_name} ({model_name}{epoch_str})\n(AUC = {roc_auc:.4f})')
        plt.legend(loc="lower right")
        plt.grid(True, alpha=0.3)
        plt.savefig(os.path.join(save_dir, f"roc_curve_{dataset_name}_{model_name}.png"), dpi=300, bbox_inches='tight')
        return roc_auc, fig

    def plot_pr_curve(self, y_true, y_probs, dataset_name="unknown", model_name="unknown", save_dir=ARTIFACTS_PATH, epochs=None):
        precision, recall, _ = precision_recall_curve(y_true, y_probs)
        pr_auc = auc(recall, precision)
        avg_precision = average_precision_score(y_true, y_probs)

        pr_data = {
            "precision": precision.tolist(),
            "recall": recall.tolist(),
            "pr_auc": float(pr_auc),
            "average_precision": float(avg_precision)
        }
        json_path = os.path.join(save_dir, f"pr_data_{dataset_name}_{model_name}.json")
        with open(json_path, 'w') as f:
            json.dump(pr_data, f, indent=2)

        fig = plt.figure(figsize=(8, 6))
        plt.plot(recall, precision, color='darkgreen', lw=2.5, label=f'{model_name} (PR-AUC = {pr_auc:.4f})')
        plt.xlabel('Recall')
        plt.ylabel('Precision')
        epoch_str = f", {epochs} Epochs" if epochs is not None else ""
        plt.title(f'Precision-Recall Curve - {dataset_name} ({model_name}{epoch_str})\nPR-AUC = {pr_auc:.4f}')
        plt.grid(True, alpha=0.3)
        plt.legend(loc='lower left')

        plot_path = os.path.join(save_dir, f"pr_curve_{dataset_name}_{model_name}.png")
        plt.savefig(plot_path, dpi=300, bbox_inches='tight')
        return pr_auc, fig

    def plot_summary_page(self, y_true, y_probs, cm, dataset_name, model_name, epochs=None):
        """Creates a single page containing ROC, PR, and Confusion Matrix."""
        # Modified layout: 2 rows. Top: ROC, PR. Bottom: CM.
        fig = plt.figure(figsize=(12, 14))
        ax_roc = plt.subplot2grid((2, 2), (0, 0))
        ax_pr = plt.subplot2grid((2, 2), (0, 1))
        ax_cm = plt.subplot2grid((2, 2), (1, 0), colspan=2)
        
        # 1. ROC Curve
        fpr, tpr, _ = roc_curve(y_true, y_probs)
        roc_auc = auc(fpr, tpr)
        ax_roc.plot(fpr, tpr, color='darkorange', lw=2, label=f'AUC = {roc_auc:.4f}')
        ax_roc.plot([0, 1], [0, 1], color='navy', lw=2, linestyle='--')
        ax_roc.set_xlim([0.0, 1.0])
        ax_roc.set_ylim([0.0, 1.05])
        ax_roc.set_xlabel('False Positive Rate')
        ax_roc.set_ylabel('True Positive Rate')
        ax_roc.set_title('ROC Curve')
        ax_roc.legend(loc="lower right")
        ax_roc.grid(True, alpha=0.3)
        
        # 2. PR Curve
        precision, recall, _ = precision_recall_curve(y_true, y_probs)
        pr_auc = auc(recall, precision)
        ax_pr.plot(recall, precision, color='darkgreen', lw=2, label=f'PR-AUC = {pr_auc:.4f}')
        ax_pr.set_xlim([0.0, 1.0])
        ax_pr.set_ylim([0.0, 1.05])
        ax_pr.set_xlabel('Recall')
        ax_pr.set_ylabel('Precision')
        ax_pr.set_title('Precision-Recall Curve')
        ax_pr.legend(loc='lower left')
        ax_pr.grid(True, alpha=0.3)
        
        # 3. Confusion Matrix
        sns.heatmap(cm, annot=True, fmt='d', cmap='Blues', 
                    xticklabels=['Normal', 'Anomaly'], yticklabels=['Normal', 'Anomaly'], 
                    ax=ax_cm, cbar=False)
        ax_cm.set_title('Confusion Matrix')
        ax_cm.set_xlabel('Predicted')
        ax_cm.set_ylabel('Actual')
        
        epoch_str = f" ({epochs} Epochs)" if epochs is not None else ""
        fig.suptitle(f"Performance Summary: {model_name} on {dataset_name}{epoch_str}", fontsize=18)
        plt.tight_layout(rect=[0, 0.03, 1, 0.95])
        return fig

    def save_training_stats(self, model_name, duration, best_acc, epochs, dataset_name, cache_dir=LOGS_PATH):
        filepath = os.path.join(cache_dir, "training_stats.json")
        all_stats = {}
        if os.path.exists(filepath):
            try:
                with open(filepath, 'r') as f:
                    all_stats = json.load(f)
            except json.JSONDecodeError:
                all_stats = {}

        current_run_key = f"{model_name}_{dataset_name}"
        all_stats[current_run_key] = {
            "model_name": model_name,
            "duration": duration,
            "best_accuracy": best_acc,
            "epochs": epochs,
            "dataset_name": dataset_name,
            "training_duration_seconds": float(duration)
        }
        with open(filepath, 'w') as f:
            json.dump(all_stats, f, indent=4)
        print(f"[INFO] Training stats saved for {dataset_name} to {filepath}")


# ==========================================
# 4. Training Pipeline
# ==========================================
def set_seed(seed: int) -> None:
    """Set global random seeds for reproducibility (PyTorch, NumPy, Python)."""
    torch.manual_seed(seed)
    np.random.seed(seed)
    random.seed(seed)
    if torch.cuda.is_available():
        torch.cuda.manual_seed_all(seed)


def train_model_pipeline(model_name, epochs=5, batch_size=32, dataset_choice='ucirvine_chest_xray',
                         training_seed=1):
    metrics = MetricsManager()
    print(f"[INFO] Training {model_name} on {device} (training seed={training_seed})...")

    # Make training reproducible
    set_seed(training_seed)

    data_root = os.path.join(DATASET_PATH, dataset_choice)
    if not os.path.isdir(data_root):
        print(f"[ERROR] Dataset subdirectory not found: {data_root}")
        print(f"        Please check if '{dataset_choice}' exists under {DATASET_PATH}")
        return

    # Ensure required split folders exist
    for split in ['train', 'test']:
        split_path = os.path.join(data_root, split)
        if not os.path.isdir(split_path):
            print(f"[ERROR] Required split folder '{split}' missing in {data_root}")
            return

    train_transform = transforms.Compose([
        transforms.Resize((IMG_SIZE, IMG_SIZE)),
        transforms.RandomHorizontalFlip(),
        transforms.ToTensor(),
        transforms.Normalize([0.485, 0.456, 0.406], [0.229, 0.224, 0.225])
    ])

    val_transform = transforms.Compose([
        transforms.Resize((IMG_SIZE, IMG_SIZE)),
        transforms.ToTensor(),
        transforms.Normalize([0.485, 0.456, 0.406], [0.229, 0.224, 0.225])
    ])

    train_dataset = datasets.ImageFolder(os.path.join(data_root, 'train'), transform=train_transform)
    val_dataset = datasets.ImageFolder(os.path.join(data_root, 'test'), transform=val_transform)

    # Generator for deterministic DataLoader shuffling when a seed is provided
    g = torch.Generator()
    g.manual_seed(training_seed)

    train_loader = DataLoader(
        train_dataset, batch_size=batch_size, shuffle=True,
        num_workers=4, pin_memory=True, generator=g
    )
    val_loader = DataLoader(
        val_dataset, batch_size=batch_size, shuffle=False,
        num_workers=4, pin_memory=True
    )

    model = build_model(model_name).to(device)
    if torch.cuda.is_available() and torch.cuda.device_count() > 1:
        print(f"[INFO] Using {torch.cuda.device_count()} GPUs for training!")
        model = nn.DataParallel(model)
        
    criterion = nn.BCEWithLogitsLoss()
    optimizer = optim.Adam(model.parameters(), lr=0.001)

    metrics.start_timer()
    best_acc = 0.0
    save_path = os.path.join(MODELS_PATH, f"{model_name}_{dataset_choice}_best.pth")

    for epoch in range(epochs):
        model.train()
        running_loss = 0.0
        for images, labels in train_loader:
            # Re-map labels: ImageFolder default is ANOMALY=0, NORMAL=1. We want NORMAL=0, ANOMALY=1.
            labels = 1 - labels 
            images, labels = images.to(device), labels.to(device).float().unsqueeze(1)
            optimizer.zero_grad()
            outputs = model(images)
            loss = criterion(outputs, labels)
            loss.backward()
            optimizer.step()
            running_loss += loss.item()

        model.eval()
        correct = total = 0
        with torch.no_grad():
            for images, labels in val_loader:
                # Re-map labels: ImageFolder default is ANOMALY=0, NORMAL=1. We want NORMAL=0, ANOMALY=1.
                labels_val = 1 - labels
                images, labels_val = images.to(device), labels_val.to(device).float().unsqueeze(1)
                outputs = model(images)
                preds = torch.sigmoid(outputs) > 0.5
                correct += (preds == labels_val).sum().item()
                total += labels_val.size(0)

        val_acc = correct / total if total > 0 else 0
        print(f"Epoch {epoch + 1}/{epochs} | Loss: {running_loss / len(train_loader):.4f} | Val Acc: {val_acc:.2%}")

        if val_acc > best_acc:
            best_acc = val_acc
            state_dict = model.module.state_dict() if isinstance(model, nn.DataParallel) else model.state_dict()
            torch.save(state_dict, save_path)

    duration = metrics.stop_timer()
    metrics.save_training_stats(model_name, duration, best_acc, epochs, dataset_choice, cache_dir=LOGS_PATH)
    return save_path, best_acc, duration


# ==========================================
# 5. Inference & Evaluation
# ==========================================
def load_trained_model(model_arch, model_path):
    model = build_model(model_arch, pretrained=False).to(device)
    model.load_state_dict(torch.load(model_path, map_location=device, weights_only=True))
    
    if torch.cuda.is_available() and torch.cuda.device_count() > 1:
        print(f"[INFO] Using {torch.cuda.device_count()} GPUs for inference!")
        model = nn.DataParallel(model)
        
    model.eval()
    return model


def predict_image(model_path, image_path, model_arch):
    model = load_trained_model(model_arch, model_path)
    transform = transforms.Compose([
        transforms.Resize((IMG_SIZE, IMG_SIZE)),
        transforms.ToTensor(),
        transforms.Normalize([0.485, 0.456, 0.406], [0.229, 0.224, 0.225])
    ])
    img = Image.open(image_path).convert('RGB')
    tensor = transform(img).unsqueeze(0).to(device)
    with torch.no_grad():
        output = model(tensor)
        prob = torch.sigmoid(output).item()
    label = "ANOMALY" if prob > 0.5 else "NORMAL"
    confidence = prob * 100 if label == "ANOMALY" else (1 - prob) * 100
    return label, confidence


def evaluate_test_set(model_path, dataset_choice='ucirvine_chest_xray', model_arch="DenseNet121"):
    # Late imports for docx to allow bootstrap installation if missing
    from docx import Document
    from docx.shared import Inches
    
    model = load_trained_model(model_arch, model_path)
    test_dir = os.path.join(DATASET_PATH, dataset_choice, 'test')

    if not os.path.isdir(test_dir):
        print(f"[ERROR] Test directory not found: {test_dir}")
        return None

    metrics_mgr = MetricsManager()
    y_true, y_probs, inference_times = [], [], []
    metrics_mgr.start_timer()

    transform = transforms.Compose([
        transforms.Resize((IMG_SIZE, IMG_SIZE)),
        transforms.ToTensor(),
        transforms.Normalize([0.485, 0.456, 0.406], [0.229, 0.224, 0.225])
    ])

    test_dataset = datasets.ImageFolder(test_dir, transform=transform)
    test_loader = DataLoader(test_dataset, batch_size=32, shuffle=False, num_workers=4, pin_memory=True)

    with torch.no_grad():
        for images, labels in test_loader:
            # Re-map labels: ImageFolder default is ANOMALY=0, NORMAL=1. We want NORMAL=0, ANOMALY=1.
            labels = 1 - labels
            
            images = images.to(device)
            start_inf = time.time()
            outputs = model(images)
            batch_inf_time = (time.time() - start_inf) * 1000
            
            probs = torch.sigmoid(outputs).cpu().numpy().flatten()
            y_true.extend(labels.numpy())
            y_probs.extend(probs)
            inference_times.extend([batch_inf_time / images.size(0)] * images.size(0))

    duration = metrics_mgr.stop_timer()
    if not y_true:
        print(f"[ERROR] No images found in {test_dir}.")
        return None

    results = metrics_mgr.calculate_metrics(y_true, y_probs)
    epochs = get_epochs_from_stats(model_arch, dataset_choice)
    roc_auc, fig_roc = metrics_mgr.plot_roc_curve(y_true, y_probs, dataset_name=dataset_choice, model_name=model_arch,
                                                  save_dir=ARTIFACTS_PATH, epochs=epochs)
    pr_auc, fig_pr = metrics_mgr.plot_pr_curve(y_true, y_probs, dataset_name=dataset_choice, model_name=model_arch,
                                               save_dir=ARTIFACTS_PATH, epochs=epochs)

    results["roc_auc"] = float(roc_auc)
    results["pr_auc"] = float(pr_auc)
    results["avg_inference_time_ms"] = np.mean(inference_times) if inference_times else 0

    print(
        f"\n{'=' * 60}\n[INFO] Evaluation complete in {duration:.1f}s\n   F1-Score: {results.get('f1_score', 0):.4f}\n   Balanced Accuracy: {results.get('balanced_accuracy', 0):.4f}\n   AUC-ROC: {roc_auc:.4f}\n   PR-AUC: {pr_auc:.4f}\n{'=' * 60}\n")

    cm = np.array(results['confusion_matrix'])
    fig_cm = metrics_mgr.plot_confusion_matrix(cm, title=f"Confusion Matrix - ({model_arch})", epochs=epochs)
    fig_cm.savefig(os.path.join(ARTIFACTS_PATH, f"confusion_matrix_{dataset_choice}_{model_arch}.png"), dpi=300,
                bbox_inches='tight')

    # Create Summary Page for PDF
    fig_summary = metrics_mgr.plot_summary_page(y_true, y_probs, cm, dataset_choice, model_arch, epochs=epochs)
    fig_summary.savefig(os.path.join(ARTIFACTS_PATH, f"summary_report_{dataset_choice}_{model_arch}.png"), dpi=300,
                        bbox_inches='tight')

    # Create PDF Report in Results folder
    pdf_path = os.path.join(RESULTS_PATH, f"evaluation_report_{dataset_choice}_{model_arch}.pdf")
    with PdfPages(pdf_path) as pdf:
        pdf.savefig(fig_roc)
        pdf.savefig(fig_pr)
        pdf.savefig(fig_cm)
        pdf.savefig(fig_summary)
    
    print(f"[SUCCESS] PDF Report saved: {pdf_path}")

    # Create DOCX Report in Results folder
    docx_path = os.path.join(RESULTS_PATH, f"evaluation_report_{dataset_choice}_{model_arch}.docx")
    doc = Document()
    doc.add_heading(f"Evaluation Report: {model_arch} on {dataset_choice}", 0)

    doc.add_paragraph(f"Model Architecture: {model_arch}")
    doc.add_paragraph(f"Dataset: {dataset_choice}")
    if epochs is not None:
        doc.add_paragraph(f"Training History: {epochs} Epochs")
    
    doc.add_paragraph(f"F1-Score: {results.get('f1_score', 0):.4f}")
    doc.add_paragraph(f"Balanced Accuracy: {results.get('balanced_accuracy', 0):.4f}")
    doc.add_paragraph(f"AUC-ROC: {roc_auc:.4f}")
    doc.add_paragraph(f"PR-AUC: {pr_auc:.4f}")

    # Add Images
    doc.add_heading('ROC Curve', level=1)
    doc.add_picture(os.path.join(ARTIFACTS_PATH, f"roc_curve_{dataset_choice}_{model_arch}.png"), width=Inches(6))
    
    doc.add_page_break()
    doc.add_heading('Precision-Recall Curve', level=1)
    doc.add_picture(os.path.join(ARTIFACTS_PATH, f"pr_curve_{dataset_choice}_{model_arch}.png"), width=Inches(6))
    
    doc.add_page_break()
    doc.add_heading('Confusion Matrix', level=1)
    doc.add_picture(os.path.join(ARTIFACTS_PATH, f"confusion_matrix_{dataset_choice}_{model_arch}.png"), width=Inches(6))
    
    doc.add_page_break()
    doc.add_heading('Performance Summary', level=1)
    doc.add_picture(os.path.join(ARTIFACTS_PATH, f"summary_report_{dataset_choice}_{model_arch}.png"), width=Inches(6))

    doc.save(docx_path)
    print(f"[SUCCESS] DOCX Report saved: {docx_path}")

    plt.close(fig_roc)
    plt.close(fig_pr)
    plt.close(fig_cm)
    plt.close(fig_summary)

    return results


def run_multiple_seeds(model_path, dataset_choice, model_arch, n_seeds=1):
    """
    Re-evaluate the same checkpoint under several random seeds.
    n_seeds controls how many independent evaluation passes are performed.
    """
    all_metrics = []
    for seed in range(n_seeds):
        set_seed(seed)
        res = evaluate_test_set(model_path, dataset_choice, model_arch)
        if res:
            all_metrics.append({
                "seed": seed,
                "f1": res.get("f1_score", 0),
                "balanced_acc": res.get("balanced_accuracy", 0),
                "mcc": res.get("mcc", 0),
                "roc_auc": res.get("roc_auc", 0),
                "pr_auc": res.get("pr_auc", 0),
                "brier": res.get("brier_score", 0),
                "avg_inference_time_ms": res.get("avg_inference_time_ms", 0)
            })

    if all_metrics:
        df = pd.DataFrame(all_metrics)
        return df
    return None


def generate_dissertation_report(models=None, dataset_name="ucirvine_chest_xray", n_seeds=1, csv_dir=RESULTS_PATH):
    if models is None:
        models = AVAILABLE_MODELS

    all_results = []
    stats_path = os.path.join(LOGS_PATH, "training_stats.json")
    all_training_stats = json.load(open(stats_path, 'r')) if os.path.exists(stats_path) else {}

    for model_name in models:
        model_path = os.path.join(MODELS_PATH, f"{model_name}_{dataset_name}_best.pth")
        if not os.path.exists(model_path): continue

        df = run_multiple_seeds(model_path, dataset_name, model_name, n_seeds=n_seeds)
        if df is not None:
            df["model"] = model_name
            df["dataset"] = dataset_name
            train_key = f"{model_name}_{dataset_name}"
            df["epochs"] = all_training_stats.get(train_key, {}).get("epochs", 0)
            df["training_time_s"] = all_training_stats.get(train_key, {}).get("training_duration_seconds", 0.0)
            all_results.append(df)

    if not all_results: return

    combined_df = pd.concat(all_results, ignore_index=True)
    combined_df["timestamp"] = time.strftime('%Y-%m-%d %H:%M:%S')
    
    # Reorder columns: identifiers first
    cols = ["timestamp", "model", "dataset", "seed"] + [
        c for c in combined_df.columns if c not in ["timestamp", "model", "dataset", "seed"]
    ]
    combined_df = combined_df[cols]
    
    csv_path = os.path.join(csv_dir, f"dissertation_report_{dataset_name}.csv")
    combined_df.to_csv(csv_path, index=False)
    print(f"\n[INFO] Full dissertation report CSV saved: {csv_path}")


# ==========================================
# 6. Chart Comparisons
# ==========================================
def get_epochs_from_stats(model_name, dataset_name, cache_dir=LOGS_PATH):
    filepath = os.path.join(cache_dir, "training_stats.json")
    if os.path.exists(filepath):
        try:
            with open(filepath, 'r') as f:
                all_stats = json.load(f)
            key = f"{model_name}_{dataset_name}"
            if key in all_stats and "epochs" in all_stats[key]:
                return all_stats[key]["epochs"]
        except json.JSONDecodeError:
            pass
    return None


def plot_comparisons(models, dataset_name, artifacts_dir=ARTIFACTS_PATH):
    overall_epochs = get_epochs_from_stats(models[0], dataset_name, cache_dir=LOGS_PATH) if models else None
    epochs_str = f" ({overall_epochs} epochs)" if overall_epochs is not None else ""

    # 1. Confusion Matrix
    num_models = len(models)
    ncols = 2
    nrows = math.ceil(num_models / ncols)

    fig, axes = plt.subplots(nrows, ncols, figsize=(6 * ncols, 6 * nrows))
    axes = axes.flatten() if nrows > 1 or ncols > 1 else [axes]

    for i, model_name in enumerate(models):
        img_path = os.path.join(artifacts_dir, f"confusion_matrix_{dataset_name}_{model_name}.png")
        if os.path.exists(img_path):
            axes[i].imshow(plt.imread(img_path))
            axes[i].axis('off')
            axes[i].set_title(f'{model_name}', fontsize=14)
        else:
            axes[i].axis('off')

    for j in range(num_models, len(axes)): fig.delaxes(axes[j])
    plt.suptitle(f'Confusion Matrix Comparison on {dataset_name}{epochs_str}', fontsize=16, y=1.02)
    plt.tight_layout(rect=[0, 0, 1, 0.98])
    plt.savefig(os.path.join(artifacts_dir, f"confusion_matrix_comparison_{dataset_name}.png"), dpi=300,
                bbox_inches='tight')
    plt.close()

    # 2. PR Curves
    fig, axes = plt.subplots(nrows, ncols, figsize=(7 * ncols, 6 * nrows))
    axes = axes.flatten() if nrows > 1 or ncols > 1 else [axes]
    colors = ['darkgreen', 'blue', 'purple', 'red', 'orange']

    for i, model_name in enumerate(models):
        json_path = os.path.join(artifacts_dir, f"pr_data_{dataset_name}_{model_name}.json")
        if os.path.exists(json_path):
            with open(json_path, 'r') as f: data = json.load(f)
            axes[i].plot(data["recall"], data["precision"], color=colors[i % len(colors)], lw=2.5,
                         label=f'{model_name} (PR-AUC = {data.get("pr_auc", 0):.4f})')
            axes[i].set_xlim([0.0, 1.0])
            axes[i].set_ylim([0.0, 1.05])
            axes[i].set_xlabel('Recall')
            axes[i].set_ylabel('Precision')
            axes[i].set_title(f'PR Curve - {model_name}')
            axes[i].grid(True, alpha=0.3)
            axes[i].legend(loc="lower left", fontsize=9)

    for j in range(num_models, len(axes)): fig.delaxes(axes[j])
    fig.suptitle(f'Precision-Recall Curves Comparison on {dataset_name}{epochs_str}', fontsize=16, y=1.02)
    plt.tight_layout(rect=[0, 0, 1, 0.98])
    plt.savefig(os.path.join(artifacts_dir, f"pr_curves_comparison_separate_{dataset_name}.png"), dpi=300,
                bbox_inches='tight')
    plt.close()

    print(f"[INFO] Comparisons plotted and saved to {artifacts_dir}")


def check_environment(dataset_choice=None, model_arch=None):
    """Performs a comprehensive check of the initial configuration, dependencies, and paths."""
    print("\n" + "="*60)
    print("INITIAL CONFIGURATION & ENVIRONMENT CHECK")
    print("="*60)
    
    # 1. Libraries and Dependencies
    print("\n[1/4] Libraries and Dependencies:")
    required = [
        "Pillow", "numpy", "matplotlib", "seaborn",
        "scikit-learn", "pandas", "boto3", "psutil",
        "python-docx", "torch", "torchvision", "torchaudio"
    ]
    all_deps_ok = True
    for pkg in required:
        try:
            ver = importlib.metadata.version(pkg)
            print(f"  [OK] {pkg:<15} : v{ver}")
        except importlib.metadata.PackageNotFoundError:
            print(f"  [MISSING] {pkg:<15} : NOT INSTALLED")
            all_deps_ok = False

    # 2. Necessary Directories
    print("\n[2/4] Necessary Directories:")
    paths = [
        ("DATASET_PATH", DATASET_PATH, False),
        ("LOGS_PATH", LOGS_PATH, True),
        ("RESULTS_PATH", RESULTS_PATH, True)
    ]
    for name, path, check_write in paths:
        exists = os.path.exists(path)
        is_dir = os.path.isdir(path) if exists else False
        writable = os.access(path, os.W_OK) if exists else False
        
        status = "[OK]" if (exists and is_dir) else "[MISSING]"
        print(f"  {status} {name:<12}: {path}")
        if exists and not is_dir:
            print(f"    ! Error: Path exists but is not a directory.")
        if exists and check_write and not writable:
            print(f"    ! Warning: Directory is not writable.")

    # 3. Dataset and Files
    if dataset_choice:
        print(f"\n[3/4] Dataset and Files ({dataset_choice}):")
        data_root = os.path.join(DATASET_PATH, dataset_choice)
        if os.path.isdir(data_root):
            print(f"  [OK] Root folder: {data_root}")
            for split in ['train', 'test']:
                split_path = os.path.join(data_root, split)
                if os.path.isdir(split_path):
                    img_count = 0
                    for r, d, files in os.walk(split_path):
                        img_count += len([f for f in files if f.lower().endswith(('.png', '.jpg', '.jpeg'))])
                    print(f"    [OK] '{split}' split found ({img_count} images).")
                else:
                    print(f"    [MISSING] '{split}' split missing.")
        else:
            print(f"  [MISSING] Dataset subdirectory not found: {data_root}")
        
        if model_arch:
            model_file = os.path.join(MODELS_PATH, f"{model_arch}_{dataset_choice}_best.pth")
            if os.path.exists(model_file):
                print(f"  [OK] Model weights: {model_file}")
            else:
                print(f"  [INFO] Model weights not found (optional for setup): {model_file}")
    else:
        print("\n[3/4] Dataset and Files: Skipped (use --dataset to check specific data)")

    # 4. Hardware Diagnostic
    print("\n[4/4] Hardware Diagnostic Report:")
    print_hardware_report()
    
    print("="*60)
    if all_deps_ok:
        print("RESULT: System configuration check complete. Ready for execution.")
    else:
        print("RESULT: Configuration incomplete. Please resolve missing dependencies.")
    print("="*60 + "\n")


# ==========================================
# 7. CLI Execution Logic
# ==========================================
if __name__ == "__main__":
    if "--mode" in sys.argv and ("check" in sys.argv or "all" in sys.argv):
        install_dependencies()
        verify_directories()
        
    initialize_logging()
    parser = argparse.ArgumentParser(description="Unified PyTorch Anomaly Detection Pipeline")
    parser.add_argument("--mode", choices=["train", "predict", "evaluate", "report", "compare", "check", "all", "listmodels", "clear"], required=True)
    parser.add_argument("--model", default="DenseNet121", choices=AVAILABLE_MODELS)
    parser.add_argument("--epochs", type=int, default=5)
    parser.add_argument("--batch", type=int, default=32)
    parser.add_argument("--dataset", default="ucirvine_chest_xray")
    parser.add_argument("--image", help="Path to image for prediction")
    parser.add_argument("--model_path", help="Path to .pth model")
    parser.add_argument("--trainingseed", type=int, default=1,
                        help="Random seed used for training (default: 1)")
    parser.add_argument("--inferenceseed", type=int, default=1,
                        help="Number of random seeds for multi-seed inference/report (default: 1)")
    args = parser.parse_args()

    # Allow overriding directories at runtime via environment variables passed as CLI arguments if needed,
    # but the primary resolution happens at module load (Section 1).

    if args.mode == "train":
        train_model_pipeline(
            args.model, args.epochs, args.batch, args.dataset,
            training_seed=args.trainingseed
        )
    elif args.mode == "predict":
        model_path = args.model_path or os.path.join(MODELS_PATH, f"{args.model}_{args.dataset}_best.pth")
        if not args.image:
            print("[ERROR] Provide --image for prediction")
        elif not os.path.exists(model_path):
            print(f"[ERROR] Model path not found: {model_path}")
        else:
            set_seed(args.inferenceseed)
            label, conf = predict_image(model_path, args.image, args.model)
            print(f"RESULT: {label} ({conf:.1f} %)")
    elif args.mode == "evaluate":
        model_path = args.model_path or os.path.join(MODELS_PATH, f"{args.model}_{args.dataset}_best.pth")
        if not os.path.exists(model_path):
            print(f"[ERROR] Model path not found: {model_path}")
            print(f"        Ensure the model exists at {model_path} or provide --model_path explicitly.")
        else:
            set_seed(args.inferenceseed)
            evaluate_test_set(model_path, args.dataset, args.model)
    elif args.mode == "report":
        generate_dissertation_report(
            dataset_name=args.dataset, n_seeds=args.inferenceseed
        )
    elif args.mode == "compare":
        models_to_compare = AVAILABLE_MODELS
        plot_comparisons(models_to_compare, args.dataset)
    elif args.mode == "check":
        check_environment(args.dataset, args.model)
    elif args.mode == "listmodels":
        list_available_models()
    elif args.mode == "clear":
        clear_directories()
    elif args.mode == "all":
        print(f"\n{'='*70}")
        print(f"RUNNING ALL-IN-ONE PIPELINE: {args.model} on {args.dataset}")
        print(f"  trainingseed={args.trainingseed}  inferenceseed={args.inferenceseed}")
        print(f"{'='*70}")

        # 1. Check
        print("\n[STEP 1/5] Environment Check")
        check_environment(args.dataset, args.model)

        # 2. Train
        print("\n[STEP 2/5] Training Model")
        train_model_pipeline(
            args.model, args.epochs, args.batch, args.dataset,
            training_seed=args.trainingseed
        )

        # 3. Evaluate
        print("\n[STEP 3/5] Evaluating Model")
        model_path = os.path.join(MODELS_PATH, f"{args.model}_{args.dataset}_best.pth")
        if os.path.exists(model_path):
            set_seed(args.inferenceseed)
            evaluate_test_set(model_path, args.dataset, args.model)
        else:
            print(f"[ERROR] Model path not found for evaluation: {model_path}")

        # 4. Compare
        print("\n[STEP 4/5] Generating Comparisons")
        models_to_compare = AVAILABLE_MODELS
        plot_comparisons(models_to_compare, args.dataset)

        # 5. Report
        print("\n[STEP 5/5] Generating Dissertation Report")
        generate_dissertation_report(
            dataset_name=args.dataset, n_seeds=args.inferenceseed
        )
        
        print(f"\n{'='*70}")
        print("[SUCCESS] Full pipeline execution completed.")
        print(f"Results are available in: {RESULTS_PATH}")
        print(f"{'='*70}\n")