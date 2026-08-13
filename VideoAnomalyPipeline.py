#!/usr/bin/env python3
"""
VideoAnomalyPipeline.py
=======================
Unified PyTorch pipeline for supervised video anomaly detection.

Mirrors the structure and features of ImageAnomalyPipeline:
  - Environment bootstrap & dependency checks
  - Path management (logs / results / artifacts / models)
  - Hardware diagnostics
  - Models: CNN-LSTM, 3D CNN (R(2+1)D), Video Transformer (VideoMAE)
  - Training with reproducible seeds
  - Evaluation (ROC, PR, confusion matrix, PDF + DOCX reports)
  - Multi-seed inference reporting for dissertation tables
  - Comparison plots across models
  - CLI modes: train | predict | evaluate | report | compare | check | all | listmodels | clear

Expected dataset layout
-----------------------
  <DATASET_PATH>/<dataset_name>/
      normal/*.mp4
      abnormal/*.mp4

Compatible with local PyCharm, Colab, and SageMaker.
"""

import sys
import subprocess
import os
import shutil
import importlib.metadata
import time

# External default dataset location (override with DATASET_DIR or SM_CHANNEL_TRAINING)
EXTERNAL_DATASET_PATH = "K:/VideoDataset"

# ==========================================
# 0. Environment Bootstrap & Dependencies
# ==========================================
def install_dependencies():
    """Install required packages if they are not already present."""
    if sys.version_info >= (3, 14):
        print(f"[ERROR] Python {sys.version_info.major}.{sys.version_info.minor} "
              "is not supported by PyTorch yet.")
        print("[ERROR] Please use Python 3.10, 3.11, or 3.12.")
        sys.exit(1)

    required_packages = [
        "Pillow", "numpy", "matplotlib", "seaborn", "scipy",
        "scikit-learn", "pandas", "boto3", "psutil",
        "python-docx", "opencv-python-headless", "transformers"
    ]

    def is_installed(package_name):
        try:
            importlib.metadata.version(package_name)
            return True
        except importlib.metadata.PackageNotFoundError:
            # Special case for OpenCV variants
            if package_name == "opencv-python-headless":
                try:
                    importlib.metadata.version("opencv-python")
                    return True
                except importlib.metadata.PackageNotFoundError:
                    pass
            return False

    missing_packages = [pkg for pkg in required_packages if not is_installed(pkg)]

    def is_nvidia_gpu_present():
        try:
            subprocess.check_output(["nvidia-smi"], stderr=subprocess.STDOUT)
            return True
        except Exception:
            return False

    gpu_present = is_nvidia_gpu_present()

    try:
        import torch
        import torchvision
        torch_ok = True
        cuda_ok = torch.cuda.is_available()
    except ImportError:
        torch_ok = False
        cuda_ok = False

    torch_packages = ["torch", "torchvision", "torchaudio"]
    missing_torch = [pkg for pkg in torch_packages if not is_installed(pkg)]

    if missing_packages:
        print(f"[INFO] Missing dependencies {missing_packages}. Installing...")
        subprocess.check_call([sys.executable, "-m", "pip", "install", *missing_packages])

    if missing_torch or (gpu_present and not cuda_ok):
        install_cmd = [sys.executable, "-m", "pip", "install"]
        if gpu_present and not cuda_ok:
            if torch_ok:
                print("[WARN] NVIDIA GPU detected but PyTorch is CPU-only. "
                      "Reinstalling with CUDA support...")
                install_cmd.extend(["torch", "torchvision", "torchaudio", "--force-reinstall"])
            else:
                print("[INFO] NVIDIA GPU detected. Installing PyTorch with CUDA support...")
                install_cmd.extend(["torch", "torchvision", "torchaudio"])
        else:
            print(f"[INFO] Missing PyTorch packages {missing_torch}. Installing...")
            install_cmd.extend(missing_torch)
        install_cmd.extend(["--index-url", "https://download.pytorch.org/whl/cu121"])
        subprocess.check_call(install_cmd)
        print("[SUCCESS] PyTorch dependencies handled successfully.\n")


# Late imports (after optional bootstrap)
import json
import math
import random
import argparse
import platform
import glob
import psutil
import numpy as np
import pandas as pd
import matplotlib.pyplot as plt
import seaborn as sns
from PIL import Image
from matplotlib.backends.backend_pdf import PdfPages

import torch
import torch.nn as nn
import torch.optim as optim
from torch.utils.data import Dataset, DataLoader, random_split
from torchvision import models
from torchvision.models import ResNet18_Weights
from torchvision.models.video import r2plus1d_18, R2Plus1D_18_Weights

try:
    from transformers import VideoMAEForVideoClassification
    TRANSFORMERS_AVAILABLE = True
except ImportError:
    TRANSFORMERS_AVAILABLE = False

try:
    import cv2
except ImportError:
    cv2 = None

from sklearn.metrics import (
    accuracy_score, precision_score, recall_score, f1_score,
    confusion_matrix, classification_report,
    roc_curve, auc,
    precision_recall_curve, average_precision_score,
    matthews_corrcoef, balanced_accuracy_score, brier_score_loss
)

# ==========================================
# 1. Configuration & Path Management
# ==========================================
BASE_DIR = os.environ.get("BASE_DIR", os.getcwd())

DEFAULT_DATASET_PATH = os.environ.get(
    "SM_CHANNEL_TRAINING",
    os.path.join(BASE_DIR, EXTERNAL_DATASET_PATH)
)
DATASET_PATH = os.environ.get("DATASET_DIR", DEFAULT_DATASET_PATH)

LOGS_PATH = os.environ.get("LOGS_DIR", os.path.join(BASE_DIR, "logs"))
os.makedirs(LOGS_PATH, exist_ok=True)

DEFAULT_RESULTS_PATH = os.environ.get("SM_MODEL_DIR", os.path.join(BASE_DIR, "results"))
RESULTS_PATH = os.environ.get("RESULTS_DIR", DEFAULT_RESULTS_PATH)
os.makedirs(RESULTS_PATH, exist_ok=True)

ARTIFACTS_PATH = os.path.join(LOGS_PATH, "artifacts")
os.makedirs(ARTIFACTS_PATH, exist_ok=True)

MODELS_PATH = os.path.join(LOGS_PATH, "models")
os.makedirs(MODELS_PATH, exist_ok=True)

# Video-specific hyper-parameters (overridable via env)
IMG_SIZE = int(os.environ.get("IMG_SIZE", 224))
SEQ_LEN = int(os.environ.get("SEQ_LEN", 16))
DEFAULT_BATCH = int(os.environ.get("BATCH_SIZE", 8))

device = torch.device("cuda" if torch.cuda.is_available() else "cpu")


def initialize_logging():
    """Redirect stdout/stderr to console + append-only master log file."""
    log_file = os.path.join(LOGS_PATH, "pipeline_execution.log")

    class Tee:
        def __init__(self, terminal, stream):
            self.terminal = terminal
            self.stream = stream

        def write(self, message):
            self.terminal.write(message)
            self.stream.write(message)
            self.stream.flush()

        def flush(self):
            self.terminal.flush()
            self.stream.flush()

        def __getattr__(self, attr):
            return getattr(self.terminal, attr)

    try:
        log_stream = open(log_file, "a", encoding="utf-8")
        log_stream.write(f"\n{'=' * 70}\n")
        log_stream.write(f"SESSION START: {time.strftime('%Y-%m-%d %H:%M:%S')}\n")
        log_stream.write(f"Command: {' '.join(sys.argv)}\n")
        log_stream.write(f"{'=' * 70}\n\n")
        sys.stdout = Tee(sys.stdout, log_stream)
        sys.stderr = Tee(sys.stderr, log_stream)
        print(f"[INFO] Session logging started (appending to {log_file})")
    except Exception as e:
        print(f"[WARN] Could not initialize log file: {e}")


def verify_directories():
    """Ensure critical directories exist and are writable."""
    print("[INFO] Path verification in progress...")

    if not os.path.exists(DATASET_PATH):
        print(f"[WARN] DATASET_PATH does not exist: {DATASET_PATH}")
        print("       (Execution may fail if the current mode requires local data)")
    elif not os.path.isdir(DATASET_PATH):
        print(f"[ERROR] DATASET_PATH exists but is not a directory: {DATASET_PATH}")
    else:
        print(f"[SUCCESS] DATASET_PATH verified: {DATASET_PATH}")

    for path_name, path_val in [("LOGS_PATH", LOGS_PATH), ("RESULTS_PATH", RESULTS_PATH)]:
        try:
            os.makedirs(path_val, exist_ok=True)
            test_file = os.path.join(path_val, ".write_test")
            with open(test_file, "w") as f:
                f.write("test")
            os.remove(test_file)
            print(f"[SUCCESS] {path_name} verified: {path_val}")
        except Exception as e:
            print(f"[ERROR] {path_name} cannot be initialised: {path_val}")
            print(f"        Details: {e}")
            sys.exit(1)


def clear_directories():
    """Delete all files in results, models, and artifacts directories."""
    targets = [
        ("Results", RESULTS_PATH),
        ("Models", MODELS_PATH),
        ("Artifacts", ARTIFACTS_PATH),
    ]
    print("\n--- Clearing Pipeline Directories ---")
    for name, path in targets:
        if os.path.exists(path):
            items_removed = 0
            for item in os.listdir(path):
                item_path = os.path.join(path, item)
                try:
                    if os.path.isfile(item_path) or os.path.islink(item_path):
                        os.unlink(item_path)
                        items_removed += 1
                    elif os.path.isdir(item_path):
                        shutil.rmtree(item_path)
                        items_removed += 1
                except Exception as e:
                    print(f"[ERROR] Failed to delete {item_path}. Reason: {e}")
            print(f"[SUCCESS] {name} directory cleared ({items_removed} items removed).")
        else:
            print(f"[INFO] {name} directory does not exist: {path}")
    print("--------------------------------------\n")


# ==========================================
# Hardware Diagnostic Report
# ==========================================
def print_hardware_report():
    print("\n--- Hardware Diagnostic Report ---")
    cpu_name = platform.processor() or "Unknown CPU Architecture"
    print(f"CPU Name: {cpu_name}")

    sys_mem = psutil.virtual_memory()
    print(f"System Memory (RAM): {sys_mem.total / (1024 ** 3):.2f} GB")

    device_type = "cuda" if torch.cuda.is_available() else "cpu"
    print(f"Target Compute Device: {device_type.upper()}")

    if torch.cuda.is_available():
        gpu_count = torch.cuda.device_count()
        print(f"Number of GPUs: {gpu_count}")
        for i in range(gpu_count):
            gpu_name = torch.cuda.get_device_name(i)
            vram_gb = torch.cuda.get_device_properties(i).total_memory / (1024 ** 3)
            print(f"  - GPU {i}: {gpu_name} ({vram_gb:.2f} GB VRAM)")
    else:
        print("Number of GPUs: 0")
        print("  - No GPU detected. Execution will fall back to CPU.")
    print("--------------------------------------\n")


# ==========================================
# 2. Model Architecture
# ==========================================
AVAILABLE_MODELS = ["CNN-LSTM", "3D CNN", "Video Transformer"]


def list_available_models():
    print("\n--- Available Video Model Architectures ---")
    for model in AVAILABLE_MODELS:
        print(f"  - {model}")
    print("--------------------------------------\n")


class CNNLSTM(nn.Module):
    """ResNet18 feature extractor + LSTM temporal head (baseline)."""

    def __init__(self, num_classes=2, lstm_hidden_size=256, lstm_layers=2):
        super().__init__()
        resnet = models.resnet18(weights=ResNet18_Weights.DEFAULT)
        self.feature_extractor = nn.Sequential(*list(resnet.children())[:-1])
        for param in self.feature_extractor.parameters():
            param.requires_grad = False
        self.lstm = nn.LSTM(
            input_size=512, hidden_size=lstm_hidden_size,
            num_layers=lstm_layers, batch_first=True
        )
        self.fc = nn.Linear(lstm_hidden_size, num_classes)

    def forward(self, x):
        # x: (B, T, C, H, W)
        b, t, c, h, w = x.size()
        cnn_out = self.feature_extractor(x.view(b * t, c, h, w))
        cnn_out = cnn_out.view(b, t, -1)
        lstm_out, _ = self.lstm(cnn_out)
        return self.fc(lstm_out[:, -1, :])


class ResNet3D(nn.Module):
    """R(2+1)D-18 spatiotemporal CNN."""

    def __init__(self, num_classes=2):
        super().__init__()
        self.model = r2plus1d_18(weights=R2Plus1D_18_Weights.DEFAULT)
        self.model.fc = nn.Linear(self.model.fc.in_features, num_classes)

    def forward(self, x):
        # x: (B, T, C, H, W) -> (B, C, T, H, W)
        return self.model(x.permute(0, 2, 1, 3, 4))


class VideoTransformer(nn.Module):
    """VideoMAE transformer fine-tuned for binary classification."""

    def __init__(self, num_classes=2):
        super().__init__()
        if not TRANSFORMERS_AVAILABLE:
            raise ImportError(
                "transformers is required for Video Transformer. "
                "Run: pip install transformers"
            )
        self.model = VideoMAEForVideoClassification.from_pretrained(
            "MCG-NJU/videomae-base-finetuned-kinetics",
            num_labels=num_classes,
            ignore_mismatched_sizes=True
        )

    def forward(self, x):
        # x: (B, T, C, H, W)
        return self.model(pixel_values=x).logits


def build_model(model_name: str, num_classes: int = 2) -> nn.Module:
    print(f"Building video model: {model_name}...")
    if model_name == "CNN-LSTM":
        return CNNLSTM(num_classes=num_classes)
    if model_name == "3D CNN":
        return ResNet3D(num_classes=num_classes)
    if model_name == "Video Transformer":
        return VideoTransformer(num_classes=num_classes)
    raise ValueError(f"Unknown model name: {model_name}")


def safe_model_name(name: str) -> str:
    return name.replace(" ", "_")


# ==========================================
# 3. Dataset
# ==========================================
class VideoFolderDataset(Dataset):
    """
    Loads fixed-length frame sequences from MP4 files.

    Layout:
        root_dir/
            normal/*.mp4   -> label 0
            abnormal/*.mp4 -> label 1
    """

    def __init__(self, root_dir: str, img_size: int = IMG_SIZE, seq_len: int = SEQ_LEN):
        if cv2 is None:
            raise ImportError("opencv-python-headless is required. pip install opencv-python-headless")

        self.img_size = img_size
        self.seq_len = seq_len
        self.video_paths = []
        self.labels = []

        for label, folder in enumerate(["normal", "abnormal"]):
            pattern = os.path.join(root_dir, folder, "*.mp4")
            files = sorted(glob.glob(pattern))
            # Also accept .avi / .mov if present
            for ext in ("*.avi", "*.mov", "*.mkv"):
                files.extend(sorted(glob.glob(os.path.join(root_dir, folder, ext))))
            self.video_paths.extend(files)
            self.labels.extend([label] * len(files))

        if not self.video_paths:
            print(f"[WARN] No video files found under {root_dir}/{{normal,abnormal}}/")

    def __len__(self):
        return len(self.video_paths)

    def __getitem__(self, idx):
        path = self.video_paths[idx]
        label = self.labels[idx]
        frames = self._read_frames(path)
        return torch.tensor(frames), torch.tensor(label, dtype=torch.long)

    def _read_frames(self, path: str) -> np.ndarray:
        cap = cv2.VideoCapture(path)
        frames = []
        while len(frames) < self.seq_len:
            ret, frame = cap.read()
            if not ret:
                break
            frame = cv2.resize(frame, (self.img_size, self.img_size))
            frame = cv2.cvtColor(frame, cv2.COLOR_BGR2RGB)
            frame = frame.astype(np.float32) / 255.0
            frame = np.transpose(frame, (2, 0, 1))  # C, H, W
            frames.append(frame)
        cap.release()

        # Pad with zeros if video is shorter than SEQ_LEN
        while len(frames) < self.seq_len:
            frames.append(np.zeros((3, self.img_size, self.img_size), dtype=np.float32))

        return np.array(frames[:self.seq_len], dtype=np.float32)


# ==========================================
# 4. Metrics & Reporting
# ==========================================
class MetricsManager:
    def __init__(self):
        self.start_time = 0.0
        self.end_time = 0.0

    def start_timer(self):
        self.start_time = time.time()

    def stop_timer(self) -> float:
        self.end_time = time.time()
        return self.end_time - self.start_time

    def calculate_metrics(self, y_true, y_pred_probs):
        y_pred = [1 if p > 0.5 else 0 for p in y_pred_probs]
        return {
            "accuracy": accuracy_score(y_true, y_pred),
            "precision": precision_score(y_true, y_pred, zero_division=0),
            "recall": recall_score(y_true, y_pred, zero_division=0),
            "f1_score": f1_score(y_true, y_pred, zero_division=0),
            "balanced_accuracy": balanced_accuracy_score(y_true, y_pred),
            "mcc": matthews_corrcoef(y_true, y_pred),
            "brier_score": brier_score_loss(y_true, y_pred_probs),
            "confusion_matrix": confusion_matrix(y_true, y_pred).tolist(),
            "report": classification_report(y_true, y_pred, output_dict=True, zero_division=0),
        }

    def plot_confusion_matrix(self, cm, labels=None, title="Confusion Matrix", epochs=None):
        if labels is None:
            labels = ["Normal", "Anomaly"]
        fig = plt.figure(figsize=(6, 4))
        sns.heatmap(cm, annot=True, fmt="d", cmap="Blues",
                    xticklabels=labels, yticklabels=labels)
        plt.xlabel("Predicted")
        plt.ylabel("Actual")
        full_title = f"{title} ({epochs} Epochs)" if epochs is not None else title
        plt.title(full_title)
        plt.tight_layout()
        return fig

    def plot_roc_curve(self, y_true, y_probs, dataset_name="unknown",
                       model_name="unknown", save_dir=ARTIFACTS_PATH, epochs=None):
        fpr, tpr, _ = roc_curve(y_true, y_probs)
        roc_auc = auc(fpr, tpr)
        safe = safe_model_name(model_name)

        with open(os.path.join(save_dir, f"roc_data_{dataset_name}_{safe}.json"), "w") as f:
            json.dump({"fpr": fpr.tolist(), "tpr": tpr.tolist(), "auc": float(roc_auc)}, f)

        fig = plt.figure(figsize=(8, 6))
        plt.plot(fpr, tpr, color="darkorange", lw=2, label=f"ROC (AUC = {roc_auc:.4f})")
        plt.plot([0, 1], [0, 1], color="navy", lw=2, linestyle="--", label="Random")
        plt.xlim([0.0, 1.0])
        plt.ylim([0.0, 1.05])
        plt.xlabel("False Positive Rate")
        plt.ylabel("True Positive Rate")
        epoch_str = f", {epochs} Epochs" if epochs is not None else ""
        plt.title(f"ROC Curve - {dataset_name} ({model_name}{epoch_str})\n(AUC = {roc_auc:.4f})")
        plt.legend(loc="lower right")
        plt.grid(True, alpha=0.3)
        plt.savefig(os.path.join(save_dir, f"roc_curve_{dataset_name}_{safe}.png"),
                    dpi=300, bbox_inches="tight")
        return roc_auc, fig

    def plot_pr_curve(self, y_true, y_probs, dataset_name="unknown",
                      model_name="unknown", save_dir=ARTIFACTS_PATH, epochs=None):
        precision, recall, _ = precision_recall_curve(y_true, y_probs)
        pr_auc = auc(recall, precision)
        avg_precision = average_precision_score(y_true, y_probs)
        safe = safe_model_name(model_name)

        with open(os.path.join(save_dir, f"pr_data_{dataset_name}_{safe}.json"), "w") as f:
            json.dump({
                "precision": precision.tolist(),
                "recall": recall.tolist(),
                "pr_auc": float(pr_auc),
                "average_precision": float(avg_precision),
            }, f, indent=2)

        fig = plt.figure(figsize=(8, 6))
        plt.plot(recall, precision, color="darkgreen", lw=2.5,
                 label=f"{model_name} (PR-AUC = {pr_auc:.4f})")
        plt.xlabel("Recall")
        plt.ylabel("Precision")
        epoch_str = f", {epochs} Epochs" if epochs is not None else ""
        plt.title(f"Precision-Recall Curve - {dataset_name} ({model_name}{epoch_str})\n"
                  f"PR-AUC = {pr_auc:.4f}")
        plt.grid(True, alpha=0.3)
        plt.legend(loc="lower left")
        plt.savefig(os.path.join(save_dir, f"pr_curve_{dataset_name}_{safe}.png"),
                    dpi=300, bbox_inches="tight")
        return pr_auc, fig

    def plot_summary_page(self, y_true, y_probs, cm, dataset_name, model_name, epochs=None):
        fig = plt.figure(figsize=(12, 14))
        ax_roc = plt.subplot2grid((2, 2), (0, 0))
        ax_pr = plt.subplot2grid((2, 2), (0, 1))
        ax_cm = plt.subplot2grid((2, 2), (1, 0), colspan=2)

        fpr, tpr, _ = roc_curve(y_true, y_probs)
        roc_auc = auc(fpr, tpr)
        ax_roc.plot(fpr, tpr, color="darkorange", lw=2, label=f"AUC = {roc_auc:.4f}")
        ax_roc.plot([0, 1], [0, 1], color="navy", lw=2, linestyle="--")
        ax_roc.set_xlim([0.0, 1.0])
        ax_roc.set_ylim([0.0, 1.05])
        ax_roc.set_xlabel("False Positive Rate")
        ax_roc.set_ylabel("True Positive Rate")
        ax_roc.set_title("ROC Curve")
        ax_roc.legend(loc="lower right")
        ax_roc.grid(True, alpha=0.3)

        precision, recall, _ = precision_recall_curve(y_true, y_probs)
        pr_auc = auc(recall, precision)
        ax_pr.plot(recall, precision, color="darkgreen", lw=2, label=f"PR-AUC = {pr_auc:.4f}")
        ax_pr.set_xlim([0.0, 1.0])
        ax_pr.set_ylim([0.0, 1.05])
        ax_pr.set_xlabel("Recall")
        ax_pr.set_ylabel("Precision")
        ax_pr.set_title("Precision-Recall Curve")
        ax_pr.legend(loc="lower left")
        ax_pr.grid(True, alpha=0.3)

        sns.heatmap(cm, annot=True, fmt="d", cmap="Blues",
                    xticklabels=["Normal", "Anomaly"],
                    yticklabels=["Normal", "Anomaly"],
                    ax=ax_cm, cbar=False)
        ax_cm.set_title("Confusion Matrix")
        ax_cm.set_xlabel("Predicted")
        ax_cm.set_ylabel("Actual")

        epoch_str = f" ({epochs} Epochs)" if epochs is not None else ""
        fig.suptitle(f"Performance Summary: {model_name} on {dataset_name}{epoch_str}",
                     fontsize=18)
        plt.tight_layout(rect=[0, 0.03, 1, 0.95])
        return fig

    def save_training_stats(self, model_name, duration, best_acc, epochs,
                            dataset_name, cache_dir=LOGS_PATH):
        filepath = os.path.join(cache_dir, "training_stats.json")
        all_stats = {}
        if os.path.exists(filepath):
            try:
                with open(filepath, "r") as f:
                    all_stats = json.load(f)
            except json.JSONDecodeError:
                all_stats = {}

        key = f"{safe_model_name(model_name)}_{dataset_name}"
        all_stats[key] = {
            "model_name": model_name,
            "duration": duration,
            "best_accuracy": best_acc,
            "epochs": epochs,
            "dataset_name": dataset_name,
            "training_duration_seconds": float(duration),
        }
        with open(filepath, "w") as f:
            json.dump(all_stats, f, indent=4)
        print(f"[INFO] Training stats saved for {dataset_name} to {filepath}")


# ==========================================
# 5. Seeding & Training
# ==========================================
def set_seed(seed: int) -> None:
    """Set global random seeds for reproducibility."""
    torch.manual_seed(seed)
    np.random.seed(seed)
    random.seed(seed)
    if torch.cuda.is_available():
        torch.cuda.manual_seed_all(seed)


def train_model_pipeline(model_name, epochs=5, batch_size=DEFAULT_BATCH,
                         dataset_choice="crime-ucf", training_seed=1):
    metrics = MetricsManager()
    print(f"[INFO] Training {model_name} on {device} (training seed={training_seed})...")
    set_seed(training_seed)

    data_root = os.path.join(DATASET_PATH, dataset_choice)
    if not os.path.isdir(data_root):
        print(f"[ERROR] Dataset subdirectory not found: {data_root}")
        print(f"        Expected layout: {data_root}/normal/*.mp4 and abnormal/*.mp4")
        return None

    full_dataset = VideoFolderDataset(data_root, img_size=IMG_SIZE, seq_len=SEQ_LEN)
    if len(full_dataset) == 0:
        print(f"[ERROR] No videos found in {data_root}")
        return None

    train_size = int(0.8 * len(full_dataset))
    val_size = len(full_dataset) - train_size
    g = torch.Generator()
    g.manual_seed(training_seed)
    train_dataset, val_dataset = random_split(
        full_dataset, [train_size, val_size], generator=g
    )

    train_loader = DataLoader(
        train_dataset, batch_size=batch_size, shuffle=True,
        num_workers=2, pin_memory=True, generator=g
    )
    val_loader = DataLoader(
        val_dataset, batch_size=batch_size, shuffle=False,
        num_workers=2, pin_memory=True
    )

    model = build_model(model_name).to(device)
    if torch.cuda.is_available() and torch.cuda.device_count() > 1:
        print(f"[INFO] Using {torch.cuda.device_count()} GPUs for training!")
        model = nn.DataParallel(model)

    criterion = nn.CrossEntropyLoss()
    optimizer = optim.Adam(model.parameters(), lr=1e-4)

    metrics.start_timer()
    best_acc = 0.0
    save_path = os.path.join(
        MODELS_PATH, f"{safe_model_name(model_name)}_{dataset_choice}_best.pth"
    )

    for epoch in range(epochs):
        model.train()
        running_loss = 0.0
        for videos, labels in train_loader:
            videos, labels = videos.to(device), labels.to(device)
            optimizer.zero_grad()
            outputs = model(videos)
            loss = criterion(outputs, labels)
            loss.backward()
            optimizer.step()
            running_loss += loss.item()

        # Validation
        model.eval()
        correct = total = 0
        with torch.no_grad():
            for videos, labels in val_loader:
                videos, labels = videos.to(device), labels.to(device)
                outputs = model(videos)
                preds = outputs.argmax(dim=1)
                correct += (preds == labels).sum().item()
                total += labels.size(0)

        val_acc = correct / total if total > 0 else 0.0
        print(f"Epoch {epoch + 1}/{epochs} | "
              f"Loss: {running_loss / max(len(train_loader), 1):.4f} | "
              f"Val Acc: {val_acc:.2%}")

        if val_acc > best_acc:
            best_acc = val_acc
            state_dict = (
                model.module.state_dict()
                if isinstance(model, nn.DataParallel)
                else model.state_dict()
            )
            torch.save(state_dict, save_path)

    duration = metrics.stop_timer()
    metrics.save_training_stats(
        model_name, duration, best_acc, epochs, dataset_choice, cache_dir=LOGS_PATH
    )
    print(f"[SUCCESS] Best model saved: {save_path} (val acc={best_acc:.2%})")
    return save_path, best_acc, duration


# ==========================================
# 6. Inference & Evaluation
# ==========================================
def load_trained_model(model_arch, model_path):
    model = build_model(model_arch).to(device)
    state = torch.load(model_path, map_location=device, weights_only=True)
    model.load_state_dict(state)
    if torch.cuda.is_available() and torch.cuda.device_count() > 1:
        print(f"[INFO] Using {torch.cuda.device_count()} GPUs for inference!")
        model = nn.DataParallel(model)
    model.eval()
    return model


def predict_video(model_path, video_path, model_arch):
    """Run a single video through a trained model. Returns (label, confidence%)."""
    if cv2 is None:
        raise ImportError("opencv-python-headless is required")

    model = load_trained_model(model_arch, model_path)
    ds = VideoFolderDataset.__new__(VideoFolderDataset)
    ds.img_size = IMG_SIZE
    ds.seq_len = SEQ_LEN
    frames = ds._read_frames(video_path)
    tensor = torch.tensor(frames).unsqueeze(0).to(device)

    with torch.no_grad():
        outputs = model(tensor)
        probs = torch.softmax(outputs, dim=1).cpu().numpy().flatten()
        anomaly_prob = float(probs[1])

    label = "ANOMALY" if anomaly_prob > 0.5 else "NORMAL"
    confidence = anomaly_prob * 100 if label == "ANOMALY" else (1 - anomaly_prob) * 100
    return label, confidence


def get_epochs_from_stats(model_name, dataset_name, cache_dir=LOGS_PATH):
    filepath = os.path.join(cache_dir, "training_stats.json")
    if os.path.exists(filepath):
        try:
            with open(filepath, "r") as f:
                all_stats = json.load(f)
            key = f"{safe_model_name(model_name)}_{dataset_name}"
            if key in all_stats and "epochs" in all_stats[key]:
                return all_stats[key]["epochs"]
        except json.JSONDecodeError:
            pass
    return None


def evaluate_test_set(model_path, dataset_choice="crime-ucf", model_arch="3D CNN"):
    from docx import Document
    from docx.shared import Inches

    model = load_trained_model(model_arch, model_path)
    data_root = os.path.join(DATASET_PATH, dataset_choice)

    if not os.path.isdir(data_root):
        print(f"[ERROR] Dataset directory not found: {data_root}")
        return None

    dataset = VideoFolderDataset(data_root, img_size=IMG_SIZE, seq_len=SEQ_LEN)
    if len(dataset) == 0:
        print(f"[ERROR] No videos found in {data_root}")
        return None

    # Use full set as evaluation (or a held-out split if preferred)
    loader = DataLoader(dataset, batch_size=DEFAULT_BATCH, shuffle=False,
                        num_workers=2, pin_memory=True)

    metrics_mgr = MetricsManager()
    y_true, y_probs, inference_times = [], [], []
    metrics_mgr.start_timer()

    with torch.no_grad():
        for videos, labels in loader:
            videos = videos.to(device)
            start_inf = time.time()
            outputs = model(videos)
            batch_ms = (time.time() - start_inf) * 1000

            probs = torch.softmax(outputs, dim=1)[:, 1].cpu().numpy()
            y_true.extend(labels.numpy().tolist())
            y_probs.extend(probs.tolist())
            per_sample = batch_ms / max(videos.size(0), 1)
            inference_times.extend([per_sample] * videos.size(0))

    duration = metrics_mgr.stop_timer()
    if not y_true:
        print("[ERROR] No predictions produced.")
        return None

    results = metrics_mgr.calculate_metrics(y_true, y_probs)
    epochs = get_epochs_from_stats(model_arch, dataset_choice)
    safe = safe_model_name(model_arch)

    roc_auc, fig_roc = metrics_mgr.plot_roc_curve(
        y_true, y_probs, dataset_name=dataset_choice, model_name=model_arch,
        save_dir=ARTIFACTS_PATH, epochs=epochs
    )
    pr_auc, fig_pr = metrics_mgr.plot_pr_curve(
        y_true, y_probs, dataset_name=dataset_choice, model_name=model_arch,
        save_dir=ARTIFACTS_PATH, epochs=epochs
    )

    results["roc_auc"] = float(roc_auc)
    results["pr_auc"] = float(pr_auc)
    results["avg_inference_time_ms"] = float(np.mean(inference_times)) if inference_times else 0.0

    print(
        f"\n{'=' * 60}\n"
        f"[INFO] Evaluation complete in {duration:.1f}s\n"
        f"   F1-Score: {results.get('f1_score', 0):.4f}\n"
        f"   Balanced Accuracy: {results.get('balanced_accuracy', 0):.4f}\n"
        f"   AUC-ROC: {roc_auc:.4f}\n"
        f"   PR-AUC: {pr_auc:.4f}\n"
        f"{'=' * 60}\n"
    )

    cm = np.array(results["confusion_matrix"])
    fig_cm = metrics_mgr.plot_confusion_matrix(
        cm, title=f"Confusion Matrix - ({model_arch})", epochs=epochs
    )
    fig_cm.savefig(
        os.path.join(ARTIFACTS_PATH, f"confusion_matrix_{dataset_choice}_{safe}.png"),
        dpi=300, bbox_inches="tight"
    )

    fig_summary = metrics_mgr.plot_summary_page(
        y_true, y_probs, cm, dataset_choice, model_arch, epochs=epochs
    )
    fig_summary.savefig(
        os.path.join(ARTIFACTS_PATH, f"summary_report_{dataset_choice}_{safe}.png"),
        dpi=300, bbox_inches="tight"
    )

    # PDF report
    pdf_path = os.path.join(
        RESULTS_PATH, f"evaluation_report_{dataset_choice}_{safe}.pdf"
    )
    with PdfPages(pdf_path) as pdf:
        pdf.savefig(fig_roc)
        pdf.savefig(fig_pr)
        pdf.savefig(fig_cm)
        pdf.savefig(fig_summary)
    print(f"[SUCCESS] PDF Report saved: {pdf_path}")

    # DOCX report
    docx_path = os.path.join(
        RESULTS_PATH, f"evaluation_report_{dataset_choice}_{safe}.docx"
    )
    doc = Document()
    doc.add_heading(f"Evaluation Report: {model_arch} on {dataset_choice}", 0)
    doc.add_paragraph(f"Model Architecture: {model_arch}")
    doc.add_paragraph(f"Dataset: {dataset_choice}")
    if epochs is not None:
        doc.add_paragraph(f"Training History: {epochs} Epochs")
    doc.add_paragraph(f"F1-Score: {results.get('f1_score', 0):.4f}")
    doc.add_paragraph(f"Balanced Accuracy: {results.get('balanced_accuracy', 0):.4f}")
    doc.add_paragraph(f"AUC-ROC: {roc_auc:.4f}")
    doc.add_paragraph(f"PR-AUC: {pr_auc:.4f}")
    doc.add_paragraph(f"Avg Inference Time: {results['avg_inference_time_ms']:.2f} ms/video")

    for heading, fname in [
        ("ROC Curve", f"roc_curve_{dataset_choice}_{safe}.png"),
        ("Precision-Recall Curve", f"pr_curve_{dataset_choice}_{safe}.png"),
        ("Confusion Matrix", f"confusion_matrix_{dataset_choice}_{safe}.png"),
        ("Performance Summary", f"summary_report_{dataset_choice}_{safe}.png"),
    ]:
        img_path = os.path.join(ARTIFACTS_PATH, fname)
        if os.path.exists(img_path):
            doc.add_page_break()
            doc.add_heading(heading, level=1)
            doc.add_picture(img_path, width=Inches(6))

    doc.save(docx_path)
    print(f"[SUCCESS] DOCX Report saved: {docx_path}")

    plt.close(fig_roc)
    plt.close(fig_pr)
    plt.close(fig_cm)
    plt.close(fig_summary)
    return results


def run_multiple_seeds(model_path, dataset_choice, model_arch, n_seeds=1):
    """Re-evaluate the same checkpoint under several random seeds."""
    all_metrics = []
    for seed in range(n_seeds):
        set_seed(seed)
        res = evaluate_test_set(model_path, dataset_choice, model_arch)
        if res:
            all_metrics.append({
                "seed": seed,
                "f1": res.get("f1_score", 0),
                "balanced_acc": res.get("balanced_accuracy", 0),
                "mcc": res.get("mcc", 0),
                "roc_auc": res.get("roc_auc", 0),
                "pr_auc": res.get("pr_auc", 0),
                "brier": res.get("brier_score", 0),
                "avg_inference_time_ms": res.get("avg_inference_time_ms", 0),
            })
    if all_metrics:
        return pd.DataFrame(all_metrics)
    return None


def generate_dissertation_report(models=None, dataset_name="crime-ucf",
                                 n_seeds=1, csv_dir=RESULTS_PATH):
    if models is None:
        models = AVAILABLE_MODELS

    all_results = []
    stats_path = os.path.join(LOGS_PATH, "training_stats.json")
    all_training_stats = (
        json.load(open(stats_path, "r")) if os.path.exists(stats_path) else {}
    )

    for model_name in models:
        model_path = os.path.join(
            MODELS_PATH, f"{safe_model_name(model_name)}_{dataset_name}_best.pth"
        )
        if not os.path.exists(model_path):
            continue

        df = run_multiple_seeds(model_path, dataset_name, model_name, n_seeds=n_seeds)
        if df is not None:
            df["model"] = model_name
            df["dataset"] = dataset_name
            train_key = f"{safe_model_name(model_name)}_{dataset_name}"
            df["epochs"] = all_training_stats.get(train_key, {}).get("epochs", 0)
            df["training_time_s"] = all_training_stats.get(
                train_key, {}
            ).get("training_duration_seconds", 0.0)
            all_results.append(df)

    if not all_results:
        print("[WARN] No trained models found for dissertation report.")
        return

    combined_df = pd.concat(all_results, ignore_index=True)
    combined_df["timestamp"] = time.strftime("%Y-%m-%d %H:%M:%S")
    cols = ["timestamp", "model", "dataset", "seed"] + [
        c for c in combined_df.columns if c not in ["timestamp", "model", "dataset", "seed"]
    ]
    combined_df = combined_df[cols]
    csv_path = os.path.join(csv_dir, f"dissertation_report_{dataset_name}.csv")
    combined_df.to_csv(csv_path, index=False)
    print(f"\n[INFO] Full dissertation report CSV saved: {csv_path}")


# ==========================================
# 7. Chart Comparisons
# ==========================================
def plot_comparisons(models, dataset_name, artifacts_dir=ARTIFACTS_PATH):
    overall_epochs = (
        get_epochs_from_stats(models[0], dataset_name) if models else None
    )
    epochs_str = f" ({overall_epochs} epochs)" if overall_epochs is not None else ""
    num_models = len(models)
    ncols = 2
    nrows = math.ceil(num_models / ncols)

    # Confusion matrices
    fig, axes = plt.subplots(nrows, ncols, figsize=(6 * ncols, 6 * nrows))
    axes = axes.flatten() if nrows > 1 or ncols > 1 else [axes]
    for i, model_name in enumerate(models):
        safe = safe_model_name(model_name)
        img_path = os.path.join(
            artifacts_dir, f"confusion_matrix_{dataset_name}_{safe}.png"
        )
        if os.path.exists(img_path):
            axes[i].imshow(plt.imread(img_path))
            axes[i].axis("off")
            axes[i].set_title(model_name, fontsize=14)
        else:
            axes[i].axis("off")
    for j in range(num_models, len(axes)):
        fig.delaxes(axes[j])
    plt.suptitle(
        f"Confusion Matrix Comparison on {dataset_name}{epochs_str}",
        fontsize=16, y=1.02
    )
    plt.tight_layout(rect=[0, 0, 1, 0.98])
    plt.savefig(
        os.path.join(artifacts_dir, f"confusion_matrix_comparison_{dataset_name}.png"),
        dpi=300, bbox_inches="tight"
    )
    plt.close()

    # PR curves
    fig, axes = plt.subplots(nrows, ncols, figsize=(7 * ncols, 6 * nrows))
    axes = axes.flatten() if nrows > 1 or ncols > 1 else [axes]
    colors = ["darkgreen", "blue", "purple", "red", "orange"]
    for i, model_name in enumerate(models):
        safe = safe_model_name(model_name)
        json_path = os.path.join(artifacts_dir, f"pr_data_{dataset_name}_{safe}.json")
        if os.path.exists(json_path):
            with open(json_path, "r") as f:
                data = json.load(f)
            axes[i].plot(
                data["recall"], data["precision"],
                color=colors[i % len(colors)], lw=2.5,
                label=f"{model_name} (PR-AUC = {data.get('pr_auc', 0):.4f})"
            )
            axes[i].set_xlim([0.0, 1.0])
            axes[i].set_ylim([0.0, 1.05])
            axes[i].set_xlabel("Recall")
            axes[i].set_ylabel("Precision")
            axes[i].set_title(f"PR Curve - {model_name}")
            axes[i].grid(True, alpha=0.3)
            axes[i].legend(loc="lower left", fontsize=9)
    for j in range(num_models, len(axes)):
        fig.delaxes(axes[j])
    fig.suptitle(
        f"Precision-Recall Curves Comparison on {dataset_name}{epochs_str}",
        fontsize=16, y=1.02
    )
    plt.tight_layout(rect=[0, 0, 1, 0.98])
    plt.savefig(
        os.path.join(artifacts_dir, f"pr_curves_comparison_separate_{dataset_name}.png"),
        dpi=300, bbox_inches="tight"
    )
    plt.close()
    print(f"[INFO] Comparisons plotted and saved to {artifacts_dir}")


# ==========================================
# 8. Environment Check
# ==========================================
def check_environment(dataset_choice=None, model_arch=None):
    print("\n" + "=" * 60)
    print("INITIAL CONFIGURATION & ENVIRONMENT CHECK")
    print("=" * 60)

    print("\n[1/4] Libraries and Dependencies:")
    required = [
        "Pillow", "numpy", "matplotlib", "seaborn", "scipy",
        "scikit-learn", "pandas", "boto3", "psutil",
        "python-docx", "opencv-python-headless", "transformers",
        "torch", "torchvision", "torchaudio"
    ]
    all_deps_ok = True
    for pkg in required:
        try:
            ver = importlib.metadata.version(pkg)
            print(f"  [OK] {pkg:<22} : v{ver}")
        except importlib.metadata.PackageNotFoundError:
            # Special case for OpenCV variants
            if pkg == "opencv-python-headless":
                try:
                    ver = importlib.metadata.version("opencv-python")
                    print(f"  [OK] opencv-python         : v{ver}")
                    continue
                except importlib.metadata.PackageNotFoundError:
                    pass
            print(f"  [MISSING] {pkg:<22} : NOT INSTALLED")
            all_deps_ok = False

    print("\n[2/4] Necessary Directories:")
    paths = [
        ("DATASET_PATH", DATASET_PATH, False),
        ("LOGS_PATH", LOGS_PATH, True),
        ("RESULTS_PATH", RESULTS_PATH, True),
    ]
    for name, path, check_write in paths:
        exists = os.path.exists(path)
        is_dir = os.path.isdir(path) if exists else False
        writable = os.access(path, os.W_OK) if exists else False
        status = "[OK]" if (exists and is_dir) else "[MISSING]"
        print(f"  {status} {name:<12}: {path}")
        if exists and not is_dir:
            print("    ! Error: Path exists but is not a directory.")
        if exists and check_write and not writable:
            print("    ! Warning: Directory is not writable.")

    if dataset_choice:
        print(f"\n[3/4] Dataset and Files ({dataset_choice}):")
        data_root = os.path.join(DATASET_PATH, dataset_choice)
        if os.path.isdir(data_root):
            print(f"  [OK] Root folder: {data_root}")
            for split in ["normal", "abnormal"]:
                split_path = os.path.join(data_root, split)
                if os.path.isdir(split_path):
                    vid_count = len([
                        f for f in os.listdir(split_path)
                        if f.lower().endswith((".mp4", ".avi", ".mov", ".mkv"))
                    ])
                    print(f"    [OK] '{split}' folder found ({vid_count} videos).")
                else:
                    print(f"    [MISSING] '{split}' folder missing.")
        else:
            print(f"  [MISSING] Dataset subdirectory not found: {data_root}")

        if model_arch:
            model_file = os.path.join(
                MODELS_PATH, f"{safe_model_name(model_arch)}_{dataset_choice}_best.pth"
            )
            if os.path.exists(model_file):
                print(f"  [OK] Model weights: {model_file}")
            else:
                print(f"  [INFO] Model weights not found (optional): {model_file}")
    else:
        print("\n[3/4] Dataset and Files: Skipped (use --dataset to check specific data)")

    print("\n[4/4] Hardware Diagnostic Report:")
    print_hardware_report()

    print("=" * 60)
    if all_deps_ok:
        print("RESULT: System configuration check complete. Ready for execution.")
    else:
        print("RESULT: Configuration incomplete. Please resolve missing dependencies.")
    print("=" * 60 + "\n")


# ==========================================
# 9. CLI Execution Logic
# ==========================================
if __name__ == "__main__":
    if "--mode" in sys.argv and ("check" in sys.argv or "all" in sys.argv):
        install_dependencies()
        verify_directories()

    initialize_logging()
    parser = argparse.ArgumentParser(
        description="Unified PyTorch Video Anomaly Detection Pipeline"
    )
    parser.add_argument(
        "--mode",
        choices=["train", "predict", "evaluate", "report", "compare",
                 "check", "all", "listmodels", "clear"],
        required=True,
    )
    parser.add_argument(
        "--model", default="3D CNN", choices=AVAILABLE_MODELS
    )
    parser.add_argument("--epochs", type=int, default=5)
    parser.add_argument("--batch", type=int, default=DEFAULT_BATCH)
    parser.add_argument("--dataset", default="crime-ucf")
    parser.add_argument("--image", help="Path to a video file for prediction "
                                        "(flag name kept for parity with image pipeline)")
    parser.add_argument("--model_path", help="Path to .pth model weights")
    parser.add_argument(
        "--trainingseed", type=int, default=1,
        help="Random seed used for training (default: 1)"
    )
    parser.add_argument(
        "--inferenceseed", type=int, default=1,
        help="Number of random seeds for multi-seed inference/report (default: 1)"
    )
    args = parser.parse_args()

    if args.mode == "train":
        train_model_pipeline(
            args.model, args.epochs, args.batch, args.dataset,
            training_seed=args.trainingseed
        )

    elif args.mode == "predict":
        model_path = args.model_path or os.path.join(
            MODELS_PATH, f"{safe_model_name(args.model)}_{args.dataset}_best.pth"
        )
        if not args.image:
            print("[ERROR] Provide --image <path_to_video> for prediction")
        elif not os.path.exists(model_path):
            print(f"[ERROR] Model path not found: {model_path}")
        else:
            set_seed(args.inferenceseed)
            label, conf = predict_video(model_path, args.image, args.model)
            print(f"RESULT: {label} ({conf:.1f} %)")

    elif args.mode == "evaluate":
        model_path = args.model_path or os.path.join(
            MODELS_PATH, f"{safe_model_name(args.model)}_{args.dataset}_best.pth"
        )
        if not os.path.exists(model_path):
            print(f"[ERROR] Model path not found: {model_path}")
            print("        Train first or provide --model_path explicitly.")
        else:
            set_seed(args.inferenceseed)
            evaluate_test_set(model_path, args.dataset, args.model)

    elif args.mode == "report":
        generate_dissertation_report(
            dataset_name=args.dataset, n_seeds=args.inferenceseed
        )

    elif args.mode == "compare":
        plot_comparisons(AVAILABLE_MODELS, args.dataset)

    elif args.mode == "check":
        check_environment(args.dataset, args.model)

    elif args.mode == "listmodels":
        list_available_models()

    elif args.mode == "clear":
        clear_directories()

    elif args.mode == "all":
        print(f"\n{'=' * 70}")
        print(f"RUNNING ALL-IN-ONE VIDEO PIPELINE: {args.model} on {args.dataset}")
        print(f"  trainingseed={args.trainingseed}  inferenceseed={args.inferenceseed}")
        print(f"{'=' * 70}")

        print("\n[STEP 1/5] Environment Check")
        check_environment(args.dataset, args.model)

        print("\n[STEP 2/5] Training Model")
        train_model_pipeline(
            args.model, args.epochs, args.batch, args.dataset,
            training_seed=args.trainingseed
        )

        print("\n[STEP 3/5] Evaluating Model")
        model_path = os.path.join(
            MODELS_PATH, f"{safe_model_name(args.model)}_{args.dataset}_best.pth"
        )
        if os.path.exists(model_path):
            set_seed(args.inferenceseed)
            evaluate_test_set(model_path, args.dataset, args.model)
        else:
            print(f"[ERROR] Model path not found for evaluation: {model_path}")

        print("\n[STEP 4/5] Generating Comparisons")
        plot_comparisons(AVAILABLE_MODELS, args.dataset)

        print("\n[STEP 5/5] Generating Dissertation Report")
        generate_dissertation_report(
            dataset_name=args.dataset, n_seeds=args.inferenceseed
        )

        print(f"\n{'=' * 70}")
        print("[SUCCESS] Full video pipeline execution completed.")
        print(f"Results are available in: {RESULTS_PATH}")
        print(f"{'=' * 70}\n")
